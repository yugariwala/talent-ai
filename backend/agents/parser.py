<<<<<<< HEAD
import os
import instructor
from groq import Groq
from pydantic import BaseModel, Field
from typing import List, Optional
import pdfplumber
import docx
from dotenv import load_dotenv

load_dotenv()

# Define the structure for parsed resume data
class Experience(BaseModel):
    title: str = Field(..., description="Job title")
    company: str = Field(..., description="Company name")
    duration: str = Field(..., description="Timeline of the job (e.g. Jan 2020 - Mar 2022)")
    description: List[str] = Field(..., description="Key responsibilities and achievements")

class Education(BaseModel):
    degree: str = Field(..., description="Degree obtained")
    institution: str = Field(..., description="University or school name")
    year: str = Field(..., description="Graduation year")

class ResumeData(BaseModel):
    name: str = Field(..., description="Full name of the candidate")
    email: Optional[str] = Field(None, description="Email address")
    phone: Optional[str] = Field(None, description="Phone number")
    summary: str = Field(..., description="A brief professional summary")
    skills: List[str] = Field(..., description="List of technical and soft skills extracted from the resume")
    experience: List[Experience] = Field(..., description="Work history")
    education: List[Education] = Field(..., description="Educational history")
    certifications: List[str] = Field(default_factory=list, description="Any certifications or awards")

class ResumeParserAgent:
    def __init__(self):
        api_key = os.getenv("GROQ_API_KEY")
        if not api_key:
            raise ValueError("GROQ_API_KEY not found in environment variables.")
        self.client = instructor.from_groq(Groq(api_key=api_key))

    def extract_text_from_pdf(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    content = page.extract_text()
                    if content:
                        text += content + "\n"
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return text

    def extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
            return ""

    def extract_text(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return self.extract_text_from_pdf(file_path)
        elif ext == ".docx":
            return self.extract_text_from_docx(file_path)
        elif ext == ".txt":
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def parse_resume(self, file_path: str) -> ResumeData:
        print(f"--- Extracting text from {os.path.basename(file_path)} ---")
        text = self.extract_text(file_path)
        
        if not text.strip():
            raise ValueError("No text could be extracted from the file.")

        print("--- Parsing structured data via Groq (Llama-3.1-70b) ---")
        try:
            resume_data = self.client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                response_model=ResumeData,
                messages=[
                    {"role": "system", "content": "You are a professional resume parser. Extract the following information accurately. If a field is missing, provide a reasonable default or empty list/null as per schema."},
                    {"role": "user", "content": f"Parse this resume:\n\n{text}"}
                ]
            )
            return resume_data
        except Exception as e:
            print(f"Error during LLM parsing: {e}")
            raise

if __name__ == "__main__":
    # Test block
    import sys
    if len(sys.argv) > 1:
        parser = ResumeParserAgent()
        result = parser.parse_resume(sys.argv[1])
        print(result.model_dump_json(indent=2))
=======
"""
Multi-Agent Talent AI — Resume Parser Agent (Agent 2)

Extracts raw text from resumes (PDF / DOCX / TXT), sends the text to
Groq API (llama-3.3-70b-versatile) with a Pydantic schema to extract
structured candidate data, assigns proficiency levels, detects red flags,
and returns a structured ParsedCandidate result dict.

Dependencies: pdfplumber, pymupdf (fitz), pytesseract, python-docx,
              groq, instructor, pydantic, Pillow

Required env vars:
    GROQ_API_KEY — Groq Cloud API key (https://console.groq.com)
"""

from __future__ import annotations

import io
import json
import logging
import statistics
from datetime import datetime, timezone
from typing import Literal

import instructor
from groq import Groq
from pydantic import BaseModel

logger = logging.getLogger(__name__)

# Model to use — llama-3.3-70b-versatile has the best JSON-mode support on Groq
GROQ_MODEL = "llama-3.3-70b-versatile"


# ---------------------------------------------------------------------------
# Step 2 — Pydantic schemas for Groq structured extraction
# ---------------------------------------------------------------------------

class RawSkillExtracted(BaseModel):
    """A single skill extracted from the resume."""
    skill_name: str
    context_sentence: str = ""
    years_mentioned: float | None = None
    proficiency: Literal["Beginner", "Intermediate", "Advanced", "Expert"]


class WorkEntryExtracted(BaseModel):
    """A single work experience entry."""
    company: str | None = None
    title: str | None = None
    start: str | None = None
    end: str | None = None
    duration_months: int | None = None
    responsibilities: list[str] = []


class EducationExtracted(BaseModel):
    """A single education entry."""
    institution: str | None = None
    degree: str | None = None
    field: str | None = None
    year: int | None = None


class CertificationExtracted(BaseModel):
    """A single certification entry."""
    name: str
    issuer: str | None = None
    year: int | None = None


class ProjectExtracted(BaseModel):
    """A single project entry."""
    name: str
    tech_stack: list[str] = []
    description: str = ""


class ParsedCandidateRaw(BaseModel):
    """Full structured extraction schema sent to Groq."""
    candidate_name: str | None = None
    email: str | None = None
    phone: str | None = None
    location: str | None = None
    work_history: list[WorkEntryExtracted] = []
    education: list[EducationExtracted] = []
    skills: list[RawSkillExtracted] = []
    certifications: list[CertificationExtracted] = []
    achievements: list[str] = []  # Only strings containing numbers, %, or $
    projects: list[ProjectExtracted] = []


# ---------------------------------------------------------------------------
# Step 1 — Text Extraction
# ---------------------------------------------------------------------------

def is_garbled(text: str) -> bool:
    """
    Returns True if extracted text is likely garbled or too short
    to be useful.

    Criteria:
    - Length < 100 chars → garbled
    - > 30% non-ASCII / non-printable characters → garbled
    """
    if len(text) < 100:
        return True
    non_ascii = sum(1 for c in text if ord(c) > 127 or not c.isprintable())
    return non_ascii / max(len(text), 1) > 0.30


def _detect_multi_column(page) -> bool:
    """
    Heuristic: checks if a pdfplumber page is multi-column by
    looking at the x0 distribution of extracted words.

    If the median x0 of the right-half words is > 60% of page
    width, we treat it as multi-column.
    """
    words = page.extract_words()
    if len(words) < 10:
        return False

    page_width = page.width
    x0_values = sorted(w["x0"] for w in words)
    mid = len(x0_values) // 2
    right_half_x0 = x0_values[mid:]

    if not right_half_x0:
        return False

    median_right = statistics.median(right_half_x0)
    return median_right > (page_width * 0.60)


def _extract_page_multi_column(page) -> str:
    """
    Re-extract text from a multi-column page by sorting words
    by (doctop, x0) reading order and joining.
    """
    words = page.extract_words()
    if not words:
        return ""

    # Sort by vertical position first, then horizontal
    sorted_words = sorted(words, key=lambda w: (w["doctop"], w["x0"]))
    return " ".join(w["text"] for w in sorted_words)


def _extract_pdf(file_path: str) -> tuple[str, bool]:
    """
    Extract text from PDF with fallback chain:
    1. pdfplumber (with multi-column detection)
    2. PyMuPDF
    3. pytesseract OCR
    """
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed. Cannot process PDF.")
        return "", False

    with pdfplumber.open(file_path) as pdf:
        pages_text = []
        for page in pdf.pages:
            # Check for multi-column layout
            if _detect_multi_column(page):
                page_text = _extract_page_multi_column(page)
            else:
                page_text = page.extract_text() or ""
            pages_text.append(page_text)
        text = "\n".join(pages_text)

    # Quality check — Fallback 1: PyMuPDF
    if is_garbled(text):
        logger.info("pdfplumber extraction garbled, falling back to PyMuPDF")
        try:
            import fitz
        except ImportError:
            logger.warning("PyMuPDF (fitz) not installed. Skipping fallback.")
            return text, False

        doc = fitz.open(file_path)
        text = "\n".join(page.get_text() for page in doc)
        doc.close()

    # Quality check — Fallback 2: pytesseract OCR
    if is_garbled(text):
        logger.info("PyMuPDF extraction garbled, falling back to OCR")
        try:
            import fitz
            import pytesseract
            from PIL import Image
            import io
        except ImportError:
            logger.warning("OCR dependencies (fitz, pytesseract, PIL) missing. Skipping OCR.")
            return text, False

        doc = fitz.open(file_path)
        ocr_pages = []
        for page in doc:
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))
            ocr_pages.append(pytesseract.image_to_string(img))
        doc.close()
        text = "\n".join(ocr_pages)
        return text, True  # ocr_processed = True

    return text, False


def _extract_docx(file_path: str) -> tuple[str, bool]:
    """
    Extract text from DOCX including paragraphs and tables.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Cannot process DOCX.")
        return "", False

    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]


    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))

    text = "\n".join(paragraphs)
    return text, False


def _extract_txt(file_path: str) -> tuple[str, bool]:
    """
    Extract text from plain text file with encoding fallback.
    """
    try:
        with open(file_path, encoding="utf-8") as f:
            text = f.read()
    except UnicodeDecodeError:
        with open(file_path, encoding="latin-1") as f:
            text = f.read()

    return text, False


def extract_text(file_path: str, file_format: str) -> tuple[str, bool]:
    """
    Extract text from a resume file.

    Args:
        file_path: Absolute path to the resume file.
        file_format: One of "pdf", "docx", "txt".

    Returns:
        Tuple of (extracted_text, ocr_processed).
    """
    fmt = file_format.lower().strip()

    if fmt == "pdf":
        return _extract_pdf(file_path)
    elif fmt == "docx":
        return _extract_docx(file_path)
    elif fmt == "txt":
        return _extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_format!r}")


# ---------------------------------------------------------------------------
# Step 3 — Groq API Call with Instructor
# ---------------------------------------------------------------------------

def extract_structured_data(
    text: str,
    criteria_rubrics: dict | None = None,
) -> ParsedCandidateRaw:
    """
    Send resume text to Groq API and get back structured candidate
    data using instructor for Pydantic enforcement.

    Uses llama-3.3-70b-versatile via Groq Cloud.
    Reads GROQ_API_KEY automatically from environment.

    Args:
        text: Raw extracted resume text.
        criteria_rubrics: Optional proficiency rubrics from the Criteria Card.

    Returns:
        ParsedCandidateRaw with all extracted fields.
    """
    client = instructor.from_groq(Groq(), mode=instructor.Mode.JSON)

    rubric_block = ""
    if criteria_rubrics:
        rubric_block = f"""
## Proficiency Rubrics (use these to assess proficiency level)
{json.dumps(criteria_rubrics, indent=2)}
"""

    prompt = f"""You are a precise resume parser. Extract ALL information from the resume below.

## Proficiency Level Guide
- Beginner: learning the skill, needs guidance, < 1 year or "familiar with" / "exposure to"
- Intermediate: works independently, 1-3 years, builds complete features
- Advanced: 3+ years, complex systems, code reviews, optimisation, team leadership
- Expert: 5+ years in that specific skill, architects systems, mentors others, makes design decisions
{rubric_block}
## Rules
- Extract EVERY skill mentioned anywhere in the resume (summary, jobs, projects, skills section)
- For each skill, estimate years from context clues in the resume
- achievements must contain a number, percentage, or dollar amount
- If a field is not present in the resume, omit it or return null
- Do not invent information not present in the resume text

## Resume Text
{text[:8000]}
"""

    for attempt in range(3):
        try:
            result = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": prompt}],
                response_model=ParsedCandidateRaw,
            )
            return result
        except Exception as e:
            logger.warning(
                "Groq extraction attempt %d/3 failed: %s", attempt + 1, e
            )
            if attempt == 2:
                # Return partial result on final failure
                logger.error("All 3 extraction attempts failed. Returning empty result.")
                return ParsedCandidateRaw()
            continue

    # Should never reach here, but safety fallback
    return ParsedCandidateRaw()


# ---------------------------------------------------------------------------
# Step 4 — Red Flag Detection
# ---------------------------------------------------------------------------

def detect_red_flags(
    parsed: ParsedCandidateRaw,
    must_haves: list[str] | None = None,
) -> list[dict]:
    """
    Detect red flags in the parsed candidate data.

    Flags detected:
    1. Must-have skills completely absent from the resume (severity: high)
    2. Job hopping — 3+ roles with tenure < 3 months (severity: medium)
    3. Employment gaps > 8 months with no stated reason (severity: medium) [from PRD §5]
    4. Advanced/Expert skills claimed but absent from work history (severity: low)

    Returns:
        List of red flag dicts with type, description, and severity.
    """
    flags: list[dict] = []

    # ── Flag 1: Must-have skills completely absent ──
    if must_haves:
        all_skill_names = [s.skill_name.lower() for s in parsed.skills]
        for mh in must_haves:
            mh_lower = mh.lower()
            if not any(
                mh_lower in skill or skill in mh_lower
                for skill in all_skill_names
            ):
                flags.append({
                    "type": "missing_must_have",
                    "description": (
                        f"Must-have requirement not found anywhere in resume: '{mh}'"
                    ),
                    "severity": "high",
                })

    # ── Flag 2: Very short tenures (< 3 months at 3+ roles) ──
    short_tenures = [
        w for w in parsed.work_history
        if w.duration_months is not None and w.duration_months < 3
    ]
    if len(short_tenures) >= 3:
        flags.append({
            "type": "job_hopping",
            "description": (
                f"{len(short_tenures)} roles with tenure under 3 months"
            ),
            "severity": "medium",
        })

    # ── Flag 3: Employment gaps > 8 months (PRD §5 — Agent 2d) ──
    # Compare end date of one role to start date of the next (sorted by start)
    dated_roles = []
    for w in parsed.work_history:
        if w.end and w.start:
            dated_roles.append(w)

    if len(dated_roles) >= 2:
        # Sort by start date string (ISO format or common date formats)
        dated_roles_sorted = sorted(
            dated_roles,
            key=lambda w: w.start or "",
        )
        for i in range(len(dated_roles_sorted) - 1):
            current_end = dated_roles_sorted[i].end
            next_start = dated_roles_sorted[i + 1].start
            # Attempt to parse dates — best-effort with common formats
            gap_months = _estimate_gap_months(current_end, next_start)
            if gap_months is not None and gap_months > 8:
                flags.append({
                    "type": "employment_gap",
                    "description": (
                        f"~{gap_months} month employment gap detected between "
                        f"'{dated_roles_sorted[i].company or 'unknown'}' and "
                        f"'{dated_roles_sorted[i + 1].company or 'unknown'}'"
                    ),
                    "severity": "medium",
                })

    # ── Flag 4: Skills claimed but absent from work history / projects ──
    work_text = " ".join(
        " ".join(w.responsibilities) + (w.title or "") + (w.company or "")
        for w in parsed.work_history
    ).lower()
    project_text = " ".join(
        p.description + " ".join(p.tech_stack)
        for p in parsed.projects
    ).lower()
    backed_text = work_text + " " + project_text

    unsubstantiated = []
    for skill in parsed.skills:
        if skill.proficiency in ("Advanced", "Expert"):
            if skill.skill_name.lower() not in backed_text:
                unsubstantiated.append(skill.skill_name)

    if len(unsubstantiated) >= 3:
        flags.append({
            "type": "unsubstantiated_skills",
            "description": (
                f"Advanced/Expert skills claimed but absent from work history: "
                f"{', '.join(unsubstantiated[:5])}"
            ),
            "severity": "low",
        })

    return flags


def _estimate_gap_months(
    end_str: str | None,
    start_str: str | None,
) -> int | None:
    """
    Best-effort estimation of the gap in months between two date strings.
    Handles common resume date formats: 'YYYY-MM', 'MM/YYYY', 'Month YYYY'.
    Returns None if parsing fails.
    """
    if not end_str or not start_str:
        return None

    from datetime import datetime as dt
    import re

    def _parse_date(s: str) -> datetime | None:
        s = s.strip()
        # Try ISO-like: 2023-06, 2023-06-15
        for fmt in ("%Y-%m", "%Y-%m-%d", "%m/%Y", "%B %Y", "%b %Y"):
            try:
                return dt.strptime(s, fmt)
            except ValueError:
                continue

        # Try extracting year-month with regex
        match = re.search(r"(\d{4})\D+(\d{1,2})", s)
        if match:
            try:
                return dt(int(match.group(1)), int(match.group(2)), 1)
            except ValueError:
                pass

        match = re.search(r"(\d{1,2})\D+(\d{4})", s)
        if match:
            try:
                return dt(int(match.group(2)), int(match.group(1)), 1)
            except ValueError:
                pass

        return None

    end_date = _parse_date(end_str)
    start_date = _parse_date(start_str)

    if not end_date or not start_date:
        return None

    delta = start_date - end_date
    months = delta.days / 30.44  # Average days per month
    return max(0, round(months))


# ---------------------------------------------------------------------------
# Step 5 — Quality Score
# ---------------------------------------------------------------------------

def compute_quality_score(parsed: ParsedCandidateRaw) -> float:
    """
    Compute what fraction of expected fields were successfully extracted.

    Checks:
    1. candidate_name present
    2. email present
    3. work_history non-empty
    4. skills non-empty
    5. education non-empty
    6. first work entry has a title
    7. achievements non-empty
    """
    checks = [
        parsed.candidate_name is not None,
        parsed.email is not None,
        len(parsed.work_history) > 0,
        len(parsed.skills) > 0,
        len(parsed.education) > 0,
        len(parsed.work_history) > 0 and parsed.work_history[0].title is not None,
        len(parsed.achievements) > 0,
    ]
    return round(sum(checks) / len(checks), 2)


# ---------------------------------------------------------------------------
# Step 6 — Main Agent Entry Point
# ---------------------------------------------------------------------------

class ParserAgent:
    """
    Resume Parser Agent (Agent 2).

    Given a file path and format, extracts text, structures it via Groq
    (llama-3.3-70b-versatile), detects red flags, and returns a dict
    ready to write to ResumeParsedData.

    Requires GROQ_API_KEY environment variable.
    """

    def __init__(self) -> None:
        # Groq client is created per extraction call to avoid stale connections.
        # The Groq SDK reads GROQ_API_KEY from the environment automatically.
        pass

    async def run(
        self,
        resume_id: str,
        file_path: str,
        file_format: str,
        must_haves: list[str] | None = None,
        criteria_rubrics: dict | None = None,
    ) -> dict:
        """
        Full parser pipeline.

        Args:
            resume_id: Unique identifier for the resume being parsed.
            file_path: Absolute path to the resume file on disk.
            file_format: One of "pdf", "docx", "txt".
            must_haves: List of must-have skills from the Criteria Card.
            criteria_rubrics: Proficiency rubrics dict from the Criteria Card.

        Returns:
            Dict with all parsed data ready to write to ResumeParsedData,
            plus orchestrator-facing metadata (hard_reject, rejection_reason,
            raw_skills).
        """
        started = datetime.now(timezone.utc)
        logger.info(
            "ParserAgent.run started | resume_id=%s file_format=%s",
            resume_id,
            file_format,
        )

        # ── Extract text ──
        text, ocr_processed = extract_text(file_path, file_format)
        logger.info(
            "Text extracted | length=%d ocr=%s", len(text), ocr_processed
        )

        # ── Extract structured data via Groq ──
        parsed = extract_structured_data(text, criteria_rubrics)
        logger.info(
            "Structured extraction complete | skills=%d work_entries=%d",
            len(parsed.skills),
            len(parsed.work_history),
        )

        # ── Detect red flags ──
        red_flags = detect_red_flags(parsed, must_haves)

        # ── Quality score ──
        quality = compute_quality_score(parsed)

        # ── Check for hard reject (high severity missing must-have) ──
        hard_reject = any(f["severity"] == "high" for f in red_flags)

        elapsed_ms = int(
            (datetime.now(timezone.utc) - started).total_seconds() * 1000
        )
        logger.info(
            "ParserAgent.run complete | resume_id=%s quality=%.2f "
            "red_flags=%d hard_reject=%s elapsed_ms=%d",
            resume_id,
            quality,
            len(red_flags),
            hard_reject,
            elapsed_ms,
        )

        return {
            "candidate_name": parsed.candidate_name,
            "contact_json": json.dumps({
                "email": parsed.email,
                "phone": parsed.phone,
                "location": parsed.location,
            }),
            "experience_json": json.dumps(
                [w.model_dump() for w in parsed.work_history]
            ),
            "skills_json": json.dumps(
                [s.model_dump() for s in parsed.skills]
            ),
            "education_json": json.dumps(
                [e.model_dump() for e in parsed.education]
            ),
            "certifications": json.dumps(
                [c.model_dump() for c in parsed.certifications]
            ),
            "achievements_json": json.dumps(parsed.achievements),
            "projects_json": json.dumps(
                [p.model_dump() for p in parsed.projects]
            ),
            "red_flags_json": json.dumps(red_flags),
            "parse_quality_score": quality,
            "ocr_processed": ocr_processed,
            "parsed_at": datetime.now(timezone.utc).isoformat(),
            # Metadata for orchestrator decisions:
            "hard_reject": hard_reject,
            "rejection_reason": next(
                (f["description"] for f in red_flags if f["severity"] == "high"),
                None,
            ),
            "raw_skills": parsed.skills,  # Passed to taxonomy agent
        }
>>>>>>> 6f21b3786fe5b50300171476c4a75a6e6958ba71
